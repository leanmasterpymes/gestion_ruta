"""Construye docs/articulo_linkedin.docx desde código.

Genera la versión LinkedIn del artículo en formato Word, con estructura
moderna por secciones y lista para copiar y pegar al editor de LinkedIn
Articles. El `.docx` contiene **solo texto**, sin imágenes, tablas, código
formateado ni fórmulas: cada artefacto visual se reemplaza por un aviso de
marcador con el patrón

    [INSERTAR AQUÍ — TIPO: <tipo> · ARCHIVO: figuras/<sub>/<nombre>.<ext> · ALT: <descripción>]

definido en la línea base de marca (Sección 4.5). El autor pega el texto en
LinkedIn y carga manualmente cada imagen donde aparece el aviso.

El artículo es ejecutivo y NO mezcla audiencias: los lectores técnicos y
estudiantes son referidos al material extendido en el repositorio (notebook
y código fuente) mediante un bloque al cierre.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


DOCX_PATH = Path(__file__).resolve().parents[1] / "docs" / "articulo_linkedin.docx"

COLOR_PRIMARIO = RGBColor(0x1A, 0x43, 0x73)
COLOR_PRIMARIO_OSCURO = RGBColor(0x0F, 0x2D, 0x52)
COLOR_ACCENTO = RGBColor(0xD9, 0x6F, 0x00)
COLOR_TEXTO = RGBColor(0x0F, 0x17, 0x29)
COLOR_TEXTO_FUERTE = RGBColor(0x1A, 0x22, 0x33)
COLOR_TEXTO_SUAVE = RGBColor(0x2D, 0x37, 0x48)
COLOR_MUDO = RGBColor(0x4F, 0x5B, 0x6F)
COLOR_AVISO = RGBColor(0xB0, 0x5C, 0x00)


# ────────────────────────────────────────────────────────────────────────────
# Helpers de formato
# ────────────────────────────────────────────────────────────────────────────


def _set_paragraph_shading(paragraph, color_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _add_horizontal_line(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C7D0DE")
    pBdr.append(bottom)
    pPr.append(pBdr)


def aviso(tipo: str, archivo: str, alt: str, numero: int | None = None) -> str:
    prefijo = f"[INSERTAR IMAGEN {numero:02d}" if numero is not None else "[INSERTAR AQUÍ"
    return f"{prefijo} — TIPO: {tipo} · ARCHIVO: {archivo} · ALT: {alt}]"


def add_titulo_principal(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = COLOR_TEXTO_FUERTE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15


def add_subtitulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = COLOR_TEXTO_SUAVE
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.45


def add_metadata(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUDO
    run.bold = True
    p.paragraph_format.space_after = Pt(16)


def add_section_heading(doc: Document, numero: str, texto: str) -> None:
    _add_horizontal_line(doc)
    p = doc.add_paragraph()
    run_num = p.add_run(f"{numero} · ")
    run_num.font.size = Pt(18)
    run_num.bold = True
    run_num.font.color.rgb = COLOR_PRIMARIO
    run_titulo = p.add_run(texto)
    run_titulo.font.size = Pt(18)
    run_titulo.bold = True
    run_titulo.font.color.rgb = COLOR_TEXTO_FUERTE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)


def add_subheading(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARIO_OSCURO
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_p(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXTO_SUAVE
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5


def add_lista(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXTO_SUAVE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.45


def add_callout(doc: Document, texto: str, color_hex: str = "E5EFFA") -> None:
    p = doc.add_paragraph()
    _set_paragraph_shading(p, color_hex)
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXTO_FUERTE
    run.italic = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)


def add_pull_quote(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_shading(p, "EEF4FB")
    run = p.add_run("“" + texto + "”")
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = COLOR_PRIMARIO_OSCURO
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6)


def add_aviso(doc: Document, texto_aviso: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_paragraph_shading(p, "FFF4E5")
    run = p.add_run(texto_aviso)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_AVISO
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)


def add_indice(doc: Document, secciones: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("EN ESTE ARTÍCULO")
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = COLOR_MUDO
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    for numero, titulo in secciones:
        p_idx = doc.add_paragraph()
        run_num = p_idx.add_run(f"  {numero}. ")
        run_num.font.size = Pt(11)
        run_num.bold = True
        run_num.font.color.rgb = COLOR_PRIMARIO
        run_titulo = p_idx.add_run(titulo)
        run_titulo.font.size = Pt(11)
        run_titulo.font.color.rgb = COLOR_TEXTO_SUAVE
        p_idx.paragraph_format.space_after = Pt(2)


def add_tech_deep_block(doc: Document) -> None:
    """Bloque que refiere a profesionales técnicos al material extendido en el repo."""
    p_titulo = doc.add_paragraph()
    _set_paragraph_shading(p_titulo, "E5EFFA")
    run = p_titulo.add_run("¿QUERÉS PROFUNDIZAR EN EL DETALLE TÉCNICO?")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARIO_OSCURO
    p_titulo.paragraph_format.space_before = Pt(14)
    p_titulo.paragraph_format.space_after = Pt(6)
    p_titulo.paragraph_format.left_indent = Cm(0.4)
    p_titulo.paragraph_format.right_indent = Cm(0.4)

    p_intro = doc.add_paragraph()
    _set_paragraph_shading(p_intro, "F4F8FD")
    run = p_intro.add_run(
        "Para profesionales técnicos, estudiantes de Investigación de Operaciones y desarrolladores "
        "que quieran entender el sistema por dentro, el repositorio incluye una versión extendida del "
        "artículo y un notebook ejecutable paso a paso con:"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXTO_FUERTE
    p_intro.paragraph_format.space_after = Pt(6)
    p_intro.paragraph_format.left_indent = Cm(0.4)
    p_intro.paragraph_format.right_indent = Cm(0.4)

    items = [
        "La formulación matemática del subproblema y la ecuación de Bellman aplicada al ruteo.",
        "La tabla de complejidad teórica y el momento exacto en que cada enfoque deja de ser viable.",
        "El código fuente completo de los siete módulos del motor.",
        "El benchmark detallado contra Google OR-Tools, con tiempos de cómputo y gap de optimalidad.",
        "Lecturas recomendadas (Held & Karp 1962, Bellman 1962, Toth & Vigo 2014).",
    ]
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph_shading(p, "F4F8FD")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXTO_FUERTE
        p.paragraph_format.space_after = Pt(2)

    p_cta = doc.add_paragraph()
    _set_paragraph_shading(p_cta, "F4F8FD")
    run = p_cta.add_run("👉 Abrir el repositorio en GitHub: github.com/leanmasterpymes/gestion_ruta")
    run.font.size = Pt(11.5)
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARIO
    p_cta.paragraph_format.space_after = Pt(10)
    p_cta.paragraph_format.left_indent = Cm(0.4)
    p_cta.paragraph_format.right_indent = Cm(0.4)


def add_cta_destacada(doc: Document, items: list[str]) -> None:
    p_titulo = doc.add_paragraph()
    _set_paragraph_shading(p_titulo, "D96F00")
    run = p_titulo.add_run("SI ESTE CONTENIDO LE RESULTÓ ÚTIL")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_titulo.paragraph_format.space_before = Pt(14)
    p_titulo.paragraph_format.space_after = Pt(8)
    p_titulo.paragraph_format.left_indent = Cm(0.4)
    p_titulo.paragraph_format.right_indent = Cm(0.4)

    for it in items:
        p = doc.add_paragraph()
        _set_paragraph_shading(p, "FFF1DC")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXTO_FUERTE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.4)


# ────────────────────────────────────────────────────────────────────────────
# Construcción del documento
# ────────────────────────────────────────────────────────────────────────────


def construir() -> Document:
    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.4)
        s.right_margin = Cm(2.4)

    # ── Portada ─────────────────────────────────────────────────────────
    add_titulo_principal(
        doc,
        "Gestión de rutas multi-sucursal con Machine Learning y programación dinámica",
    )
    add_subtitulo(
        doc,
        "Una arquitectura abierta, reproducible y de bajo costo para que cualquier distribuidora con "
        "varios centros de distribución deje de planificar rutas como islas y empiece a operar un plan "
        "diario coordinado.",
    )
    add_metadata(doc, "LECTURA, 7–9 MIN  ·  OPEN SOURCE, MIT  ·  PYMES, LOGÍSTICA  ·  CASO DE ESTUDIO REPRODUCIBLE")

    add_indice(doc, [
        ("1", "El problema operativo"),
        ("2", "La arquitectura propuesta"),
        ("3", "Predecir antes de planificar"),
        ("4", "Cómo se optimiza la ruta"),
        ("5", "Coordinación entre centros"),
        ("6", "Validación contra el estándar industrial"),
        ("7", "Cinco beneficios operativos medibles"),
        ("8", "Cómo probar el sistema"),
    ])

    # ── Demos ───────────────────────────────────────────────────────────
    add_section_heading(doc, "▣", "Demo y código sin instalación")
    add_p(doc, "Antes de leer, abre el sistema funcionando en tu navegador:")
    add_lista(doc, [
        "⭐ Repositorio en GitHub (código abierto, MIT): github.com/leanmasterpymes/gestion_ruta. Clónalo y ejecuta el sistema completo en tu computadora.",
        "Notebook en Google Colab: colab.research.google.com/github/leanmasterpymes/gestion_ruta/blob/main/notebooks/01_ruteo_multinivel.ipynb. El flujo completo paso a paso, sin instalar nada.",
        "👉 Planificador interactivo en Streamlit (sin instalación): https://gestionruta-fvzdvkcjthzfavm9fcmk9c.streamlit.app/. Plan diario con mapa interactivo, métricas del modelo de demanda, comparativa antes/después y descarga del plan en CSV.",
    ])

    # ── Sección 1 ───────────────────────────────────────────────────────
    add_section_heading(doc, "1", "El problema operativo que se repite en la mayoría de las distribuidoras")
    add_p(
        doc,
        "En la mayoría de las distribuidoras con varios centros de distribución que he visitado "
        "(sucursales, depósitos regionales o incluso fábricas), dos camiones de la misma empresa "
        "terminan a tres cuadras del mismo cliente el mismo día. Cada centro planifica su ruta como "
        "una isla: arma su propia plantilla, reparte sus pedidos, y nadie se entera de que el centro "
        "vecino está pasando por la misma calle dos horas después. El resultado se paga en silencio: "
        "kilómetros duplicados, combustible perdido, horas extra del chofer y el clásico camión flash "
        "subcontratado a última hora cuando la demanda real superó la capacidad asignada.",
    )
    add_p(doc, "El patrón suele ser el mismo:")
    add_lista(doc, [
        "Cada centro estima su demanda al ojo sobre el promedio de las últimas semanas, sin un modelo que capture estacionalidad ni tendencia.",
        "La asignación cliente–centro se hereda desde hace años y nadie la cuestiona, aun cuando dos centros hayan crecido en zonas que ya se solapan.",
        "Cuando el plan del día queda corto, se subcontrata un camión flash con un costo 30 a 50% superior al de un camión propio. Nadie mide cuánto se va al año por este renglón.",
        "Cuando el plan queda holgado, los choferes regresan a la rampa con espacio vacío y se justifican horas que no se trabajaron en ruta. La utilización real de la flota nunca aparece en el reporte mensual.",
    ])
    add_callout(
        doc,
        "Para responder a esta situación desarrollé un sistema de código abierto, disponible en GitHub: "
        "github.com/leanmasterpymes/gestion_ruta, que aborda el problema desde la predicción de demanda "
        "hasta el plan diario coordinado entre los centros de distribución.",
    )

    # ── Sección 2 ───────────────────────────────────────────────────────
    add_section_heading(doc, "2", "La arquitectura propuesta")
    add_p(
        doc,
        "La solución se compone de cuatro capas, todas basadas en software libre y reproducibles desde "
        "el repositorio. El sistema toma dos insumos principales: el histórico de pedidos de los "
        "clientes y la disponibilidad de cada centro de distribución (cantidad de camiones y capacidad "
        "de cada uno). Conocer la disponibilidad antes de asignar es lo que permite respetar la "
        "capacidad real de despacho de cada centro y evitar la sobrecarga.",
    )
    add_aviso(doc, aviso(
        "diagrama",
        "figuras/diagramas/01_arquitectura_sistema.png",
        "Pipeline del sistema con dos inputs (histórico de pedidos y disponibilidad por centro) que alimentan la predicción de demanda, el clustering de clientes, la programación dinámica, el plan de rutas coordinado y el dashboard.",
        numero=1,
    ))
    add_lista(doc, [
        "Predicción de demanda por cliente: un modelo aprende los patrones de pedido de cada cliente y proyecta cuánto pedirá mañana, capturando estacionalidad semanal, tendencia mensual y comportamiento individual.",
        "Clustering inteligente de clientes: los clientes se agrupan en zonas de ruteo balanceadas por la capacidad real del camión, no por divisiones administrativas heredadas.",
        "Optimización de la ruta: cada zona se resuelve con un algoritmo clásico de optimización que minimiza kilometraje y respeta ventanas horarias.",
        "Coordinación entre centros: antes de salir, el sistema reasigna clientes fronterizos para que dos centros no se pisen el territorio.",
    ])
    add_p(
        doc,
        "El resultado es un plan diario donde cada cliente tiene un centro de distribución, una ruta y "
        "una posición en la secuencia, con cifras reproducibles a partir de un caso de estudio que "
        "cualquiera puede correr en su computadora.",
    )

    # ── Sección 3 ───────────────────────────────────────────────────────
    add_section_heading(doc, "3", "Predecir antes de planificar")
    add_p(
        doc,
        "El insumo principal del ruteo es la demanda esperada de cada cliente para el día siguiente. "
        "Sin esa predicción, todo el resto del sistema opera con promedios planos que no capturan la "
        "realidad: lunes y martes pico, domingo valle, clientes que crecen, clientes que cambian de "
        "hábitos.",
    )
    add_p(doc, "El modelo aprende de cuatro tipos de señales:")
    add_lista(doc, [
        "El calendario (día de la semana, mes, semana del año), que captura los ciclos predecibles del negocio.",
        "El historial reciente del cliente (cuánto pidió hace 1, 7 y 14 días), que captura la inercia de pedidos y los hábitos individuales.",
        "La tendencia (medias móviles de 7 y 30 días), que suaviza ruido y muestra si el cliente está creciendo, estable o decreciendo.",
        "El perfil del cliente (tipo de comercio, ventana horaria, nivel de demanda base).",
    ])
    add_p(
        doc,
        "La validación se hace de forma rigurosa: los últimos 30 días del histórico se reservan como "
        "prueba, sin que el modelo los vea durante el aprendizaje. Esto evita la fuga de información "
        "típica de los splits aleatorios sobre series de tiempo, donde un modelo puede \"ver el futuro\" "
        "durante el entrenamiento y reportar métricas optimistas que después no se reproducen en "
        "producción.",
    )
    add_p(
        doc,
        "Sobre el caso de estudio del repositorio (50 clientes, 540 días de histórico), el modelo logra "
        "un error de 10.1% en la demanda total esperada por día (cifra de clase mundial para "
        "planificación de capacidad) y un error promedio de 11.3 unidades por cliente y día. La métrica "
        "que más importa para la planificación es el error agregado por día: si la flota total va a "
        "recibir alrededor de 870 unidades en pedidos para mañana, el modelo se equivoca en promedio un "
        "10%, suficiente para dimensionar correctamente cuántos camiones cargar.",
    )
    add_aviso(doc, aviso(
        "imagen",
        "figuras/imagenes/02_mapa_clientes_clusters.png",
        "Mapa de los 50 clientes coloreados por centro de distribución asignado, con los 5 centros marcados como cuadrados.",
        numero=2,
    ))

    # ── Sección 4 ───────────────────────────────────────────────────────
    add_section_heading(doc, "4", "Cómo se optimiza la ruta")
    add_p(
        doc,
        "La optimización de rutas se apoya en una idea de los años sesenta que sigue siendo el corazón "
        "de los mejores sistemas industriales: descomponer un problema grande en piezas pequeñas, "
        "resolver cada pieza una sola vez y reutilizar el resultado. La técnica se llama programación "
        "dinámica, y el algoritmo clásico que la materializa para el ruteo es Held-Karp (1962).",
    )
    add_pull_quote(
        doc,
        "¿Cuál es el camino más corto para visitar este conjunto de clientes y volver al centro de "
        "distribución? Esa es la pregunta. La programación dinámica la responde sin probar todas las "
        "combinaciones posibles.",
    )
    add_p(
        doc,
        "El método tiene un límite práctico: cuando una ruta crece de 15 clientes en adelante, el "
        "cálculo se vuelve demasiado pesado. La estrategia del sistema es respetar ese límite y usar "
        "descomposición por zonas: el clustering reduce 50 clientes a varias rutas pequeñas, cada una "
        "con menos de 15 clientes, donde el algoritmo exacto resuelve el óptimo en milisegundos. Cuando "
        "una zona excede los 15 clientes, se sustituye automáticamente por una variante aproximada que "
        "sigue siendo óptima en la práctica.",
    )
    add_p(
        doc,
        "El resultado, en cifras del propio sistema, es contundente: un plan completo de 50 clientes y "
        "5 centros de distribución se calcula en menos de 10 milisegundos.",
    )

    # ── Sección 5 ───────────────────────────────────────────────────────
    add_section_heading(doc, "5", "Coordinación entre centros, el verdadero cambio de paradigma")
    add_p(
        doc,
        "El gran cambio que introduce esta arquitectura no es el algoritmo de optimización en sí, sino "
        "el paso previo de asignar clientes a centros de distribución con visión global. La asignación "
        "tradicional, que es la que opera hoy en la mayoría de las distribuidoras, asigna cada cliente "
        "al centro más cercano y termina ahí. Funciona razonablemente bien cuando la red es pequeña, "
        "pero cuando crece, dos centros empiezan a invadirse el territorio.",
    )
    add_p(
        doc,
        "La asignación coordinada plantea el problema desde otro lugar: cada centro tiene una capacidad "
        "real (cantidad de camiones multiplicada por la capacidad de cada camión), y cada cliente debe "
        "asignarse de modo que se minimice la distancia total respetando esa capacidad. El sistema "
        "resuelve este planteo en milisegundos y elimina los solapamientos sin necesidad de redibujar "
        "zonas a mano.",
    )
    add_aviso(doc, aviso(
        "imagen",
        "figuras/imagenes/03_plan_rutas_resuelto.png",
        "Mapa con las rutas finales resueltas, líneas coloreadas saliendo de cada centro de distribución hacia sus clientes en secuencia óptima.",
        numero=3,
    ))

    # ── Sección 6 ───────────────────────────────────────────────────────
    add_section_heading(doc, "6", "Validación contra el estándar industrial")
    add_p(
        doc,
        "Para validar la calidad del sistema, el motor propio se compara lado a lado con Google "
        "OR-Tools, el estándar industrial de problemas de ruteo. Sobre el mismo caso de estudio:",
    )
    add_aviso(doc, aviso(
        "tabla",
        "figuras/tablas/04_benchmark_ortools.png",
        "Tabla comparativa: ruta única 12 clientes, diferencia 0%, sistema propio 50× más rápido. Plan completo 50 clientes, diferencia +15.7% bajo restricción real de centros múltiples.",
        numero=4,
    ))
    add_callout(
        doc,
        "En el caso pequeño, el sistema propio encuentra el óptimo idéntico al estándar industrial, "
        "50 veces más rápido. En el caso grande, OR-Tools logra menor kilometraje porque consolida toda "
        "la flota en un único depósito ficticio (un esquema irrealizable en la operación real, donde "
        "cada camión sale de su centro). Bajo la misma restricción, el sistema propio queda dentro "
        "de un margen razonable del óptimo industrial.",
    )

    # ── Sección 7 ───────────────────────────────────────────────────────
    add_section_heading(doc, "7", "Cinco beneficios operativos medibles")
    add_p(
        doc,
        "Las cifras siguientes corresponden a referencias públicas del sector (Gartner, McKinsey, "
        "Project44, FarEye, ORTEC, OptimoRoute) cuando una arquitectura de planificación con predicción "
        "de demanda y ruteo automatizado se implementa con disciplina. No son promesas: son rangos "
        "documentados de mejora con su traducción directa al impacto en el negocio.",
    )
    add_subheading(doc, "1.  15 a 30% menos kilometraje total")
    add_p(
        doc,
        "Menos combustible, menos desgaste de flota, menos peajes y menos tiempo improductivo en ruta. "
        "Si una distribuidora corre 50.000 km al mes, esto se traduce en un ahorro de 7.500 a 15.000 km "
        "mensuales.",
    )
    add_subheading(doc, "2.  15 a 25% más clientes por ruta")
    add_p(
        doc,
        "Mayor productividad por chofer y por camión. Si hoy un camión visita 20 clientes al día, "
        "puede llegar a 23 o 25 con el mismo recurso. Cuando la red crece, no es necesario contratar "
        "más camiones de inmediato.",
    )
    add_subheading(doc, "3.  20 a 25% menos costo de combustible")
    add_p(
        doc,
        "Ahorro directo en el renglón más visible del estado de resultados de logística. Si el "
        "combustible representa entre 25 y 30% del costo logístico total, esto se traduce en una "
        "reducción de 5 a 7% en el costo logístico total.",
    )
    add_subheading(doc, "4.  10 a 20% menos horas extra del chofer")
    add_p(
        doc,
        "Menos pago de horas extra, menor rotación de personal y menor riesgo laboral. El chofer "
        "fatigado tiene entre 2 y 3 veces más probabilidad de accidente; estabilizar la jornada "
        "protege a la operación y al negocio.",
    )
    add_subheading(doc, "5.  30 a 60% menos subcontrataciones flash")
    add_p(
        doc,
        "El ahorro más subestimado de toda la arquitectura. Un camión flash subcontratado cuesta "
        "entre 30 y 50% más que un camión propio. Si hoy 10% de las rutas se cubren con flash y se "
        "baja a 4 o 5%, ese es ahorro puro al cierre del mes, mes tras mes.",
    )
    add_aviso(doc, aviso(
        "imagen",
        "figuras/imagenes/05_antes_despues.png",
        "Comparativa antes y después en tres barras: kilometraje total, número de rutas y costo estimado.",
        numero=5,
    ))

    # ── Sección 8 ───────────────────────────────────────────────────────
    add_section_heading(doc, "8", "Cómo probar el sistema")
    add_subheading(doc, "Opción rápida, sin instalar nada")
    add_p(
        doc,
        "Hace clic en el badge \"Open in Colab\" del repositorio. Las celdas ejecutan en orden el flujo "
        "completo: caso de estudio, predicción, clustering, optimización, validación contra el "
        "estándar industrial y mapa del plan resuelto. En menos de dos minutos, la solución está "
        "corriendo en una pestaña del navegador.",
    )
    add_subheading(doc, "Opción para implementarlo, clonar el repositorio")
    add_p(
        doc,
        "🔗 El código completo está publicado bajo licencia MIT en github.com/leanmasterpymes/gestion_ruta. "
        "Cualquiera puede clonarlo, correr el sistema en su computadora y conectar sus propios datos de "
        "clientes. El planificador abre en el navegador y permite cargar un archivo de clientes propio "
        "para generar un plan personalizado.",
    )

    # ── Bloque tech-deep (referencia al material extendido) ────────────
    add_tech_deep_block(doc)

    # ── CTA ────────────────────────────────────────────────────────────
    add_cta_destacada(doc, [
        "👀 Abra el notebook en Colab antes de cerrar la pestaña, corre con un solo clic, sin instalación.",
        "⭐ Marque el repositorio en GitHub para darle visibilidad y poder clonarlo en su empresa.",
        "💬 Comparta el artículo con su responsable de logística o de operaciones; el sistema puede evaluarse en una sola jornada con los datos de su propia distribuidora.",
        "📩 Si desea adaptar la solución a su flota o integrarla a sus sistemas, queda abierto el canal de mensajes directos.",
    ])

    # ── Bio del autor ──────────────────────────────────────────────────
    _add_horizontal_line(doc)
    p_nombre = doc.add_paragraph()
    run_n = p_nombre.add_run("Manuel Antonio Pérez Ogando")
    run_n.font.size = Pt(13)
    run_n.bold = True
    run_n.font.color.rgb = COLOR_TEXTO_FUERTE
    p_nombre.paragraph_format.space_after = Pt(2)

    p_role = doc.add_paragraph()
    run_r = p_role.add_run(
        "Ingeniero industrial · MSc en Gestión Estratégica para el Desarrollo de Software · Profesor "
        "de Investigación de Operaciones"
    )
    run_r.font.size = Pt(10.5)
    run_r.font.color.rgb = COLOR_TEXTO_SUAVE
    p_role.paragraph_format.space_after = Pt(2)

    p_cred = doc.add_paragraph()
    run_c = p_cred.add_run(
        "Especialista en mapeo, análisis y mejora de procesos · Certificado Lean Six Sigma Green Belt "
        "(LSSGB) · Certificado en Power BI · Estudiante de Matemática Aplicada, mención Informática "
        "en la UASD."
    )
    run_c.font.size = Pt(10)
    run_c.font.color.rgb = COLOR_MUDO
    p_cred.paragraph_format.space_after = Pt(8)

    p_serie = doc.add_paragraph()
    run_s = p_serie.add_run(
        "Leanmaster Pymes · Entrega de la serie semanal sobre ciencia de datos aplicada a la "
        "productividad empresarial."
    )
    run_s.font.size = Pt(10)
    run_s.italic = True
    run_s.font.color.rgb = COLOR_MUDO

    return doc


def main() -> None:
    doc = construir()
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"DOCX escrito en: {DOCX_PATH}")
    print(f"Tamaño: {DOCX_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
